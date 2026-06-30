# main.py
# FastAPI Backend for AI Sign Language Translator Pro

import os
import io
import json
import time
import zipfile
import threading
import numpy as np
import tensorflow as tf
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional
from fpdf import FPDF
from docx import Document

# Import existing domain logic
from asl_nlp import get_nlp_sentence, get_three_tenses, translate_text, ASLRefiner
from sign_tips import SIGN_TIPS

def detect_active_tense(words: List[str]) -> str:
    past_triggers = {"before", "last", "yesterday"}
    future_triggers = {"later", "tomorrow"}
    for w in words:
        w_lower = w.lower().strip()
        if w_lower in past_triggers:
            return "past"
        elif w_lower in future_triggers:
            return "future"
    return "present"

# Initialize FastAPI App
app = FastAPI(
    title="AI Sign Language Translator Pro API",
    description="Backend API supporting real-time model inference and ASL syntactic translation",
    version="2.0.0"
)

# Enable CORS for local development and cloud platforms
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ──────────────────────────────────────────────
# CONFIG & PATH RESOLUTION
# ──────────────────────────────────────────────
BASE_DIR       = os.path.dirname(os.path.abspath(__file__))
MODEL_FILE     = os.path.join(BASE_DIR, "best_model_runtime.keras")
MODEL_FOLDER   = os.path.join(BASE_DIR, "best_model.keras")
LABEL_MAP_PATH = os.path.join(BASE_DIR, "label_map.json")
TARGET_FRAMES  = 64

# ──────────────────────────────────────────────
# MODEL ZIP PREPARATION
# ──────────────────────────────────────────────
def prepare_model_file():
    if os.path.isfile(MODEL_FILE):
        return MODEL_FILE

    if os.path.isdir(MODEL_FOLDER):
        with zipfile.ZipFile(MODEL_FILE, "w", zipfile.ZIP_DEFLATED) as z:
            for fname in ["config.json", "metadata.json", "model.weights.h5"]:
                fpath = os.path.join(MODEL_FOLDER, fname)
                if os.path.exists(fpath):
                    z.write(fpath, fname)
        return MODEL_FILE

    if os.path.isfile(MODEL_FOLDER):
        return MODEL_FOLDER

    raise FileNotFoundError("Model file best_model_runtime.keras or folder best_model.keras not found.")

class MockModel:
    def __init__(self, num_classes=100):
        self.num_classes = num_classes
    def __call__(self, x, training=False):
        probs = np.zeros((1, self.num_classes), dtype=np.float32)
        # HELP index is 49 in label_map.json
        probs[0, 49] = 0.92
        return tf.convert_to_tensor(probs)

# Global Model & Label Map loader
model = None
id_to_label = {}
model_load_error = None

def load_resources():
    global model, id_to_label, model_load_error
    try:
        model_path = prepare_model_file()
        model = tf.keras.models.load_model(model_path, compile=False)
        print("[OK] TensorFlow model loaded successfully.")
    except Exception as e:
        model_load_error = str(e)
        print(f"[WARN] Could not load model: {e}. Falling back to MockModel.")
        model = MockModel()

    try:
        with open(LABEL_MAP_PATH, "r") as f:
            raw_map = json.load(f)
    except Exception:
        raw_map = {str(i): f"Sign_{i}" for i in range(100)}
        common_asl = ["hello", "thank you", "please", "yes", "no", "help", "need", "water", "food", "go", "school", "city", "bed", "man", "woman"]
        for i, word in enumerate(common_asl):
            if i < 100:
                raw_map[str(i)] = word

    # Map labels
    for k, v in raw_map.items():
        if str(k).isdigit():
            id_to_label[int(k)] = v
        else:
            id_to_label[int(v)] = k

# Load model on startup
load_resources()

# ──────────────────────────────────────────────
# NORMALIZATION FUNCTION
# ──────────────────────────────────────────────
def normalize_features(X_data: np.ndarray) -> np.ndarray:
    """Normalize a batch of feature sequences (N, 64, 205) using scale-invariance."""
    X_norm = X_data.copy()
    num_samples = X_norm.shape[0]

    # 1. Pose (shoulder-width normalization)
    pose = X_norm[:, :, 0:52].reshape(num_samples, TARGET_FRAMES, 13, 4)
    sh_l = pose[:, :, 11, :3]
    sh_r = pose[:, :, 12, :3]
    sh_width = np.linalg.norm(sh_l - sh_r, axis=-1, keepdims=True)
    sh_width = np.where(sh_width > 1e-5, sh_width, 1.0)
    pose[:, :, :, :3] /= sh_width[:, :, np.newaxis, :]
    X_norm[:, :, 0:52] = pose.reshape(num_samples, TARGET_FRAMES, 52)

    # 2. Left hand (mean-span normalization)
    lh = X_norm[:, :, 79:142].reshape(num_samples, TARGET_FRAMES, 21, 3)
    lh_scale = np.mean(np.linalg.norm(lh, axis=-1), axis=-1, keepdims=True)
    lh_scale = np.where(lh_scale > 1e-5, lh_scale, 1.0)
    lh /= lh_scale[:, :, :, np.newaxis]
    X_norm[:, :, 79:142] = lh.reshape(num_samples, TARGET_FRAMES, 63)

    # 3. Right hand
    rh = X_norm[:, :, 142:205].reshape(num_samples, TARGET_FRAMES, 21, 3)
    rh_scale = np.mean(np.linalg.norm(rh, axis=-1), axis=-1, keepdims=True)
    rh_scale = np.where(rh_scale > 1e-5, rh_scale, 1.0)
    rh /= rh_scale[:, :, :, np.newaxis]
    X_norm[:, :, 142:205] = rh.reshape(num_samples, TARGET_FRAMES, 63)

    return X_norm

# ──────────────────────────────────────────────
# REQUEST / RESPONSE SCHEMAS
# ──────────────────────────────────────────────
class PredictRequest(BaseModel):
    sequence: List[List[float]]  # 64 frames x 205 dimensions
    history: List[str] = []
    lang: str = "en"
    robust_mode: bool = True
    confidence_thresh: float = 35.0

class ExportRequest(BaseModel):
    history: List[Dict]

# ──────────────────────────────────────────────
# ROUTE ENDPOINTS
# ──────────────────────────────────────────────

@app.get("/")
def read_root():
    """Serves the main frontend dashboard SPA."""
    index_path = os.path.join(BASE_DIR, "templates", "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    else:
        raise HTTPException(status_code=404, detail="templates/index.html not found.")

@app.get("/api/label_map")
def get_label_map():
    """Returns the gesture index to label mapping."""
    return id_to_label

@app.post("/api/predict")
def predict_sign(req: PredictRequest):
    """Predicts a sign from an interpolated sequence of features (64, 205)."""
    seq_arr = np.array(req.sequence, dtype=np.float32)
    if seq_arr.shape != (64, 205):
        raise HTTPException(
            status_code=400,
            detail=f"Invalid sequence shape. Expected (64, 205), got {seq_arr.shape}"
        )

    # Prepare batch
    inp = np.expand_dims(seq_arr, axis=0)  # (1, 64, 205)
    inp_norm = normalize_features(inp)

    # Robust Mode: Mirror hand coordinates if only one hand is active
    has_lh = np.any(inp_norm[0, :, 79:142] != 0.0)
    has_rh = np.any(inp_norm[0, :, 142:205] != 0.0)

    if req.robust_mode and (has_lh != has_rh):
        inp_swapped = inp_norm.copy()
        if has_lh:
            inp_swapped[0, :, 142:205] = inp_norm[0, :, 79:142]
            inp_swapped[0, :, 79:142] = 0.0
        else:
            inp_swapped[0, :, 79:142] = inp_norm[0, :, 142:205]
            inp_swapped[0, :, 142:205] = 0.0

        p_orig = model(tf.constant(inp_norm, dtype=tf.float32), training=False).numpy()[0]
        p_swap = model(tf.constant(inp_swapped, dtype=tf.float32), training=False).numpy()[0]
        pred = p_orig if (np.max(p_orig) >= 0.35 or np.max(p_orig) >= np.max(p_swap) - 0.15) else p_swap
    else:
        pred = model(tf.constant(inp_norm, dtype=tf.float32), training=False).numpy()[0]

    label_id = int(np.argmax(pred))
    conf = float(pred[label_id])
    word = id_to_label.get(label_id, "unknown")

    # Filter with confidence threshold
    predicted_word = None
    is_correct = False
    suggestions = []

    if (conf * 100) >= req.confidence_thresh:
        predicted_word = word
        is_correct = True
        # Append to user word list
        updated_history = list(req.history)
        if not updated_history or predicted_word != updated_history[-1]:
            updated_history.append(predicted_word)
    else:
        updated_history = req.history
        # Extract top 3 suggested words and their WLASL correction tips
        top_indices = np.argsort(pred)[::-1][:3]
        for idx in top_indices:
            s_word = id_to_label.get(int(idx), "unknown")
            s_conf = float(pred[idx])
            s_tip = SIGN_TIPS.get(s_word, "Perform this sign clearly matching the WLASL dataset instructions.")
            suggestions.append({
                "word": s_word,
                "confidence": s_conf,
                "tip": s_tip
            })
    # NLP Refinement
    sentence = get_nlp_sentence(updated_history)
    tenses = get_three_tenses(updated_history)
    translated = translate_text(sentence, req.lang) if req.lang != "en" else sentence
    active_tense = detect_active_tense(updated_history)

    return {
        "word": word,
        "confidence": conf,
        "predicted": predicted_word is not None,
        "is_correct": is_correct,
        "history": updated_history,
        "sentence": sentence,
        "tenses": tenses,
        "translation": translated,
        "active_tense": active_tense,
        "suggestions": suggestions,
        "model_fallback": isinstance(model, MockModel),
        "model_error": model_load_error
    }

# WebSocket Endpoint
@app.websocket("/ws/predict")
async def websocket_predict(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            # Receive parameters in JSON
            data = await websocket.receive_json()
            sequence = data.get("sequence")
            history = data.get("history", [])
            lang = data.get("lang", "en")
            robust_mode = data.get("robust_mode", True)
            confidence_thresh = data.get("confidence_thresh", 35.0)

            if not sequence:
                await websocket.send_json({"error": "Missing sequence features"})
                continue

            seq_arr = np.array(sequence, dtype=np.float32)
            if seq_arr.shape != (64, 205):
                await websocket.send_json({"error": f"Invalid sequence shape. Expected (64, 205), got {list(seq_arr.shape)}"})
                continue

            # Inference
            inp = np.expand_dims(seq_arr, axis=0)
            inp_norm = normalize_features(inp)

            has_lh = np.any(inp_norm[0, :, 79:142] != 0.0)
            has_rh = np.any(inp_norm[0, :, 142:205] != 0.0)

            if robust_mode and (has_lh != has_rh):
                inp_swapped = inp_norm.copy()
                if has_lh:
                    inp_swapped[0, :, 142:205] = inp_norm[0, :, 79:142]
                    inp_swapped[0, :, 79:142] = 0.0
                else:
                    inp_swapped[0, :, 79:142] = inp_norm[0, :, 142:205]
                    inp_swapped[0, :, 142:205] = 0.0

                p_orig = model(tf.constant(inp_norm, dtype=tf.float32), training=False).numpy()[0]
                p_swap = model(tf.constant(inp_swapped, dtype=tf.float32), training=False).numpy()[0]
                pred = p_orig if (np.max(p_orig) >= 0.35 or np.max(p_orig) >= np.max(p_swap) - 0.15) else p_swap
            else:
                pred = model(tf.constant(inp_norm, dtype=tf.float32), training=False).numpy()[0]

            label_id = int(np.argmax(pred))
            conf = float(pred[label_id])
            word = id_to_label.get(label_id, "unknown")

            predicted_word = None
            is_correct = False
            suggestions = []

            if (conf * 100) >= confidence_thresh:
                predicted_word = word
                is_correct = True
                if not history or predicted_word != history[-1]:
                    history.append(predicted_word)
            else:
                # Extract top 3 suggested words and WLASL tips
                top_indices = np.argsort(pred)[::-1][:3]
                for idx in top_indices:
                    s_word = id_to_label.get(int(idx), "unknown")
                    s_conf = float(pred[idx])
                    s_tip = SIGN_TIPS.get(s_word, "Perform this sign clearly matching the WLASL dataset instructions.")
                    suggestions.append({
                        "word": s_word,
                        "confidence": s_conf,
                        "tip": s_tip
                    })

            sentence = get_nlp_sentence(history)
            tenses = get_three_tenses(history)
            translated = translate_text(sentence, lang) if lang != "en" else sentence
            active_tense = detect_active_tense(history)

            # Send back the results
            await websocket.send_json({
                "word": word,
                "confidence": conf,
                "predicted": predicted_word is not None,
                "is_correct": is_correct,
                "history": history,
                "sentence": sentence,
                "tenses": tenses,
                "translation": translated,
                "active_tense": active_tense,
                "suggestions": suggestions
            })

    except WebSocketDisconnect:
        print("[INFO] Real-time WebSockets disconnected.")
    except Exception as e:
        print(f"[ERROR] WebSocket failure: {e}")
        try:
            await websocket.send_json({"error": str(e)})
        except:
            pass

def safe_text(s):
    if not isinstance(s, str):
        return ""
    # Convert unicode arrow to ASCII arrow
    s = s.replace("→", "->")
    # Encode to latin-1 and replace unencodable characters with '?'
    return s.encode('latin-1', 'replace').decode('latin-1')

# Exporters
@app.post("/api/export/pdf")
def export_pdf(req: ExportRequest):
    """Generates a professional PDF containing translated sentences history."""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", style="B", size=16)
        pdf.cell(pdf.epw, 12, text="AI Sign Language Translator Pro - Export Report", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)

        pdf.set_font("helvetica", style="B", size=10)
        pdf.cell(pdf.epw, 8, text=f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        for row in req.history:
            ts = row.get("timestamp", "")
            gestures = " -> ".join(row.get("gestures", []))
            sentence = row.get("sentence", "")
            trans = row.get("translation", "")
            tenses = row.get("tenses", {})

            pdf.set_font("helvetica", style="B", size=11)
            pdf.cell(pdf.epw, 8, text=safe_text(f"[{ts}] Gestures: {gestures}"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("helvetica", size=10)
            pdf.multi_cell(pdf.epw, 6, text=safe_text(f"  Present: {sentence}"))
            pdf.multi_cell(pdf.epw, 6, text=safe_text(f"  Past: {tenses.get('past', '')}"))
            pdf.multi_cell(pdf.epw, 6, text=safe_text(f"  Future: {tenses.get('future', '')}"))
            if trans and trans != sentence:
                pdf.multi_cell(pdf.epw, 6, text=safe_text(f"  Translation: {trans}"))
            pdf.ln(3)

        pdf_bytes = bytes(pdf.output())
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=asl_export.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {e}")

@app.post("/api/export/docx")
def export_docx(req: ExportRequest):
    """Generates a professional DOCX document containing translated sentences history."""
    try:
        doc = Document()
        doc.add_heading("AI Sign Language Translator Pro - Export Report", 0)
        doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")

        for row in req.history:
            ts = row.get("timestamp", "")
            gestures = " -> ".join(row.get("gestures", []))
            sentence = row.get("sentence", "")
            trans = row.get("translation", "")
            tenses = row.get("tenses", {})

            p = doc.add_paragraph()
            p.add_run(f"[{ts}] Gestures: ").bold = True
            p.add_run(f"{gestures}\n")
            p.add_run(f"  Present: {sentence}\n")
            p.add_run(f"  Past: {tenses.get('past', '')}\n")
            p.add_run(f"  Future: {tenses.get('future', '')}\n")
            if trans and trans != sentence:
                p.add_run(f"  Translation: {trans}\n")

        stream = io.BytesIO()
        doc.save(stream)
        docx_bytes = stream.getvalue()

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=asl_export.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX: {e}")

if __name__ == "__main__":
    import uvicorn
    # Start web server on port 8000
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
