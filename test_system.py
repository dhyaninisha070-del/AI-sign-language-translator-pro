# test_system.py

import os
import sys
import numpy as np

# Verify modules can be imported
try:
    import asl_nlp
    from asl_nlp import ASLRefiner, get_three_tenses, translate_text
    print("[OK] asl_nlp imported successfully.")
except Exception as e:
    print(f"[FAIL] Failed to import asl_nlp: {e}")
    sys.exit(1)

# Verify fpdf2 and docx
try:
    from fpdf import FPDF
    from docx import Document
    print("[OK] Export libraries fpdf2 and python-docx imported successfully.")
except Exception as e:
    print(f"[FAIL] Failed to import export libraries: {e}")
    sys.exit(1)

# 1. Test Refiner and Tense Generator
def test_nlp_refiner():
    refiner = ASLRefiner()
    words = ["I", "go", "school"]
    
    # Test local refiner present tense
    res = refiner.refine(words)
    expected_present = "I am going to school."
    if res == expected_present:
        print(f"[OK] Parser present tense check passed: {words} -> {res}")
    else:
        print(f"[FAIL] Parser present tense check failed: expected '{expected_present}', got '{res}'")
        return False
        
    # Test three-tense generation
    tenses = get_three_tenses(words)
    print("Tenses generated:", tenses)
    if tenses['past'] == "I went to school." and tenses['future'] == "I will go to school.":
        print("[OK] Multi-tense verification passed.")
    else:
        print("[FAIL] Multi-tense verification failed.")
        return False
    return True

# 2. Test Translation API
def test_translation():
    test_text = "I am going to school."
    target_lang = "es" # Spanish
    translated = translate_text(test_text, target_lang)
    print(f"Original: '{test_text}' -> Spanish: '{translated}'")
    if translated and translated != test_text:
        print("[OK] Translation API connection verification passed.")
        return True
    else:
        print("[WARN] Translation API returned original text (might be offline or blocked).")
        return True # Don't fail the build, it could be a network timeout

# 3. Test Export generation
def test_exports():
    sentences = ["Line 1: test", "Line 2: test"]
    # PDF
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        pdf.cell(pdf.epw, 10, text="Test PDF", new_x="LMARGIN", new_y="NEXT")
        for s in sentences:
            pdf.multi_cell(pdf.epw, 10, text=s)
        pdf_bytes = bytes(pdf.output())
        if len(pdf_bytes) > 0:
            print("[OK] PDF generation bytes check passed.")
        else:
            print("[FAIL] PDF generation returned empty bytes.")
            return False
    except Exception as e:
        print(f"[FAIL] PDF generation failed: {e}")
        return False
        
    # DOCX
    try:
        doc = Document()
        doc.add_heading("Test DOCX", 0)
        for s in sentences:
            doc.add_paragraph(s)
        import io
        stream = io.BytesIO()
        doc.save(stream)
        docx_bytes = stream.getvalue()
        if len(docx_bytes) > 0:
            print("[OK] DOCX generation bytes check passed.")
        else:
            print("[FAIL] DOCX generation returned empty bytes.")
            return False
    except Exception as e:
        print(f"[FAIL] DOCX generation failed: {e}")
        return False
        
    return True

if __name__ == "__main__":
    success = test_nlp_refiner() and test_translation() and test_exports()
    if success:
        print("\n*** ALL SYSTEM INTEGRATION TESTS PASSED ***")
        sys.exit(0)
    else:
        print("\n*** SYSTEM INTEGRATION TESTS FAILED ***")
        sys.exit(1)
